import os
import re
from pdf2image import convert_from_path
import pytesseract
from docx import Document
from PIL import Image
from config import UPLOAD_FOLDER, RESULT_FOLDER, TESSERACT_CMD, OCR_LANG

# Définir le chemin de Tesseract (si nécessaire)
pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD

def pdf_to_images(pdf_path, dpi=300):
    """Convertit un PDF en une liste d'images PIL."""
    images = convert_from_path(pdf_path, dpi=dpi)
    return images

def image_to_text(image):
    """Extrait le texte d'une image via tesseract OCR."""
    text = pytesseract.image_to_string(image, lang=OCR_LANG)
    return text

def normalize_arabic_text(text):
    """Nettoyage et normalisation du texte arabe extrait."""
    text = text.replace("\x0c", " ")  # caractères de fin de page
    text = re.sub(r'[ \t]+', ' ', text)  # espaces multiples
    text = re.sub(r'\n{3,}', '\n\n', text)  # retours multiples
    return text.strip()

def assemble_docx(text_blocks, out_path):
    """Assemble des blocs de texte dans un fichier DOCX."""
    doc = Document()
    for block in text_blocks:
        paragraph = doc.add_paragraph()
        paragraph.add_run(block)
    doc.save(out_path)
    return out_path
